"""Real-world Word constructs the synthetic fixture's dialect missed (R8/§6.3:
each was identified as a coverage gap of the reference extractor and is pinned
here as a minimal by-construction fixture, built in-test with python-docx).

Covers:
- localized style names (German Word: "Überschrift 1", "Beschriftung", …) —
  detection must run on locale-independent signals (outlineLvl, SEQ fields,
  numPr), not English style names;
- inline images inside text paragraphs (text must survive; N drawings = N figures);
- caption paragraphs *above* their figure;
- nested tables (table inside a cell);
- adjacent distinct lists (ordered then bulleted must not merge);
- ordered-ness derived from the numbering part, not the style name;
- comparator: None scalar counts (producer doesn't extract) must not flag;
  non-gating metrics excluded from `fired_flags`.
"""

from __future__ import annotations

import binascii
import io
import struct
import zlib
from pathlib import Path

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx_parse_eval import config as C
from docx_parse_eval.adapters import ooxml_reference
from docx_parse_eval.comparator import _count_metric, compare, fired_flags
from docx_parse_eval.io import read_record

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _png(width: int = 4, height: int = 4) -> bytes:
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", binascii.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 0, 0, 0, 0)
    raw = b"".join(b"\x00" + b"\x80" * width for _ in range(height))
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw, 9)) + chunk(b"IEND", b"")


def _extract(doc: docx.Document, tmp_path):
    path = tmp_path / "case.docx"
    doc.save(path)
    return ooxml_reference.extract(path)


# --- localized styles ---------------------------------------------------------
def test_localized_heading_style_detected_via_outline_level(tmp_path):
    doc = docx.Document()
    st = doc.styles.add_style("Überschrift Eins", WD_STYLE_TYPE.PARAGRAPH)
    ppr = st.element.get_or_add_pPr()
    lvl = OxmlElement("w:outlineLvl")
    lvl.set(qn("w:val"), "0")
    ppr.append(lvl)
    doc.add_paragraph("Einleitung", style="Überschrift Eins")
    doc.add_paragraph("Fließtext über das Lager.")

    rec = _extract(doc, tmp_path)
    assert [(h.text, h.level) for h in rec.headings] == [("Einleitung", 1)]
    assert rec.element_sequence == ["heading", "paragraph"]


def test_localized_caption_detected_via_seq_field(tmp_path):
    doc = docx.Document()
    doc.add_picture(io.BytesIO(_png()))
    # What Word's Insert-Caption emits in a German install: a "Beschriftung"-
    # styled paragraph whose number is a SEQ field. No English style name.
    doc.styles.add_style("Beschriftung", WD_STYLE_TYPE.PARAGRAPH)
    p = doc.add_paragraph(style="Beschriftung")
    p.add_run("Abbildung ")
    r = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    r._r.append(begin)
    r = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = r" SEQ Abbildung \* ARABIC "
    r._r.append(instr)
    r = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r._r.append(end)
    p.add_run("1: Förderband-Layout")

    rec = _extract(doc, tmp_path)
    assert rec.element_sequence == ["figure", "caption"]
    assert rec.figures[0].has_caption
    assert "Förderband-Layout" in rec.figures[0].caption_text


# --- figures ---------------------------------------------------------------------
def test_inline_image_keeps_paragraph_text_and_counts_each_drawing(tmp_path):
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("Der Aufbau misst 42 mm, siehe ")
    p.add_run().add_picture(io.BytesIO(_png()))
    p.add_run(" und ")
    p.add_run().add_picture(io.BytesIO(_png()))

    rec = _extract(doc, tmp_path)
    assert len(rec.figures) == 2  # N drawings = N figures, not 1
    assert "42 mm" in (rec.full_text_normalized or "")  # text not swallowed
    assert rec.element_sequence == ["paragraph", "figure", "figure"]


def test_caption_above_figure_is_associated(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Figure 1: Rack elevation.", style="Caption")
    doc.add_picture(io.BytesIO(_png()))

    rec = _extract(doc, tmp_path)
    assert rec.figures[0].has_caption
    assert rec.figures[0].caption_text == "Figure 1: Rack elevation."


# --- lists ---------------------------------------------------------------------
def test_adjacent_ordered_and_bulleted_lists_stay_distinct(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Receive", style="List Number")
    doc.add_paragraph("Store", style="List Number")
    doc.add_paragraph("Point A", style="List Bullet")
    doc.add_paragraph("Point B", style="List Bullet")

    rec = _extract(doc, tmp_path)
    assert [(lst.n_items, lst.is_ordered) for lst in rec.lists] == [(2, True), (2, False)]


def test_orderedness_comes_from_numbering_part_not_style_name(tmp_path):
    # "List Number" resolves to a decimal numFmt, "List Bullet" to bullet —
    # via the numbering part, which survives localization; the old
    # '"Number" in style-name' heuristic does not.
    doc = docx.Document()
    doc.add_paragraph("Eins", style="List Number")
    rec = _extract(doc, tmp_path)
    assert rec.lists[0].is_ordered is True

    doc = docx.Document()
    doc.add_paragraph("Punkt", style="List Bullet")
    rec = _extract(doc, tmp_path)
    assert rec.lists[0].is_ordered is False


def test_direct_numpr_on_normal_style_is_a_list(tmp_path):
    # Toolbar-created lists carry numPr as direct formatting on an ordinary
    # style ("List Paragraph"/localized equivalents) — no "List *" name.
    doc = docx.Document()
    doc.add_paragraph("plain lead-in")
    p = doc.add_paragraph("Erster Punkt")
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), "5")  # default template: decimal
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.append(numpr)

    rec = _extract(doc, tmp_path)
    assert [(lst.n_items, lst.is_ordered) for lst in rec.lists] == [(1, True)]


# --- nested tables ---------------------------------------------------------------
def test_nested_table_is_emitted_and_its_text_conserved(tmp_path):
    doc = docx.Document()
    outer = doc.add_table(rows=2, cols=2)
    outer.cell(0, 1).text = "outer-text"
    inner = outer.cell(0, 0).add_table(rows=1, cols=2)
    inner.cell(0, 0).text = "nested-a"
    inner.cell(0, 1).text = "nested-b"

    rec = _extract(doc, tmp_path)
    assert len(rec.tables) == 2  # outer + nested, like tree-shaped parsers
    assert rec.element_sequence.count("table") == 2
    assert "nested-a" in (rec.full_text_normalized or "")  # invisible to cell.text


# --- comparator semantics ----------------------------------------------------------
def test_none_scalar_count_never_flags():
    # None = producer does not extract the feature (schema 0.2) ≠ 0 found.
    r = _count_metric("hyperlink_count", 7, None)
    assert not r.flag and r.prediction_value == "not-extracted"
    assert _count_metric("footnote_count", None, 3).flag is False
    assert _count_metric("hyperlink_count", 7, 0).flag is True  # real loss still fires


def test_non_gating_metrics_excluded_from_fired_flags():
    gold = read_record(FIXTURES / "synthetic.gold.json")
    pred = gold.model_copy(deep=True)
    pred.tables[0].has_header = not pred.tables[0].has_header
    results = compare(gold, pred)
    assert "table_header_detection" in fired_flags(results, include_non_gating=True)
    assert "table_header_detection" not in fired_flags(results)
    assert "table_header_detection" in C.NON_GATING_METRICS


def _docling_doc(texts=(), tables=(), pictures=(), groups=(), body_children=()):
    return {
        "texts": list(texts),
        "tables": list(tables),
        "pictures": list(pictures),
        "groups": list(groups),
        "body": {"children": [{"$ref": r} for r in body_children]},
    }


def test_docling_notes_layer_excluded():
    # Comment/annotation text (content_layer "notes") is not document content.
    from docx_parse_eval.adapters.docling_adapter import extract_from_dict

    rec = extract_from_dict(_docling_doc(
        texts=[
            {"self_ref": "#/texts/0", "label": "text", "text": "body text"},
            {"self_ref": "#/texts/1", "label": "text", "text": "a comment",
             "content_layer": "notes"},
        ],
        body_children=["#/texts/0", "#/texts/1"],
    ))
    assert rec.word_count == 2
    assert "comment" not in (rec.full_text_normalized or "")


def test_docling_enumerated_items_make_list_ordered():
    from docx_parse_eval.adapters.docling_adapter import extract_from_dict

    rec = extract_from_dict(_docling_doc(
        texts=[
            {"self_ref": "#/texts/0", "label": "list_item", "text": "one", "enumerated": True},
            {"self_ref": "#/texts/1", "label": "list_item", "text": "two", "enumerated": True},
        ],
        groups=[{"self_ref": "#/groups/0", "label": "list", "name": "list",
                 "children": [{"$ref": "#/texts/0"}, {"$ref": "#/texts/1"}]}],
        body_children=["#/groups/0"],
    ))
    assert [(lst.n_items, lst.is_ordered) for lst in rec.lists] == [(2, True)]


def test_docling_list_item_children_are_not_dropped():
    # Definition bodies etc. nest under list items (list_after_num_headers).
    from docx_parse_eval.adapters.docling_adapter import extract_from_dict

    rec = extract_from_dict(_docling_doc(
        texts=[
            {"self_ref": "#/texts/0", "label": "list_item", "text": "item",
             "children": [{"$ref": "#/groups/1"}]},
            {"self_ref": "#/texts/1", "label": "text", "text": "Term 1"},
        ],
        groups=[
            {"self_ref": "#/groups/0", "label": "list", "name": "list",
             "children": [{"$ref": "#/texts/0"}]},
            {"self_ref": "#/groups/1", "label": "inline", "name": "group",
             "children": [{"$ref": "#/texts/1"}]},
        ],
        body_children=["#/groups/0"],
    ))
    assert "Term 1" in (rec.full_text_normalized or "")


def test_docling_table_children_text_not_double_counted_but_pictures_kept():
    # Rich-cell content is in table_cells AND in the table's children tree;
    # only pictures/nested tables may be emitted from the subtree.
    from docx_parse_eval.adapters.docling_adapter import extract_from_dict

    rec = extract_from_dict(_docling_doc(
        texts=[{"self_ref": "#/texts/0", "label": "text", "text": "cell words"}],
        pictures=[{"self_ref": "#/pictures/0", "label": "picture", "captions": []}],
        tables=[{
            "self_ref": "#/tables/0", "label": "table",
            "data": {"num_rows": 1, "num_cols": 1,
                     "table_cells": [{"text": "cell words",
                                      "start_row_offset_idx": 0, "start_col_offset_idx": 0}]},
            "children": [{"$ref": "#/groups/0"}],
        }],
        groups=[{"self_ref": "#/groups/0", "label": "inline", "name": "group",
                 "children": [{"$ref": "#/texts/0"}, {"$ref": "#/pictures/0"}]}],
        body_children=["#/tables/0"],
    ))
    assert rec.word_count == 2  # "cell words" once, not twice
    assert len(rec.figures) == 1  # the in-cell picture still surfaces


# --- OOXML: content controls, cell figures, Title ----------------------------------
def test_inline_sdt_text_is_extracted(tmp_path):
    # Form fields / checkbox labels wrap runs in inline w:sdt — invisible to
    # python-docx Paragraph.text (docx_checkboxes gap).
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("Choose ")
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "as many as you like"
    r.append(t)
    content.append(r)
    sdt.append(content)
    p._p.append(sdt)

    rec = _extract(doc, tmp_path)
    assert "as many as you like" in (rec.full_text_normalized or "")


def test_figure_inside_table_cell_is_counted(tmp_path):
    doc = docx.Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "label"
    t.cell(0, 1).paragraphs[0].add_run().add_picture(io.BytesIO(_png()))

    rec = _extract(doc, tmp_path)
    assert len(rec.figures) == 1
    assert rec.element_sequence == ["table", "figure"]


def test_title_style_is_a_level1_heading(tmp_path):
    doc = docx.Document()
    doc.add_paragraph("Warehouse Spec", style="Title")
    doc.add_heading("Scope", level=1)

    rec = _extract(doc, tmp_path)
    assert [h.level for h in rec.headings] == [1, 1]


# --- figure = raster image; shapes are a separate diagnostic -----------------------
def test_vml_shape_without_imagedata_is_a_shape_not_a_figure(tmp_path):
    # docx_vml_images: a positioned VML shape (watermark) has no v:imagedata —
    # it is not an image, and counting it flagged Docling's CORRECT count.
    from docx.oxml import parse_xml

    doc = docx.Document()
    p = doc.add_paragraph()
    pict = parse_xml(
        '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:v="urn:schemas-microsoft-com:vml">'
        '<v:shape style="position:absolute;width:100pt;height:50pt"/></w:pict>'
    )
    r = OxmlElement("w:r")
    r.append(pict)
    p._p.append(r)
    doc.add_paragraph("body text")

    rec = _extract(doc, tmp_path)
    assert len(rec.figures) == 0
    assert rec.vector_shape_count == 1


def test_grouped_drawing_counts_each_raster_image(tmp_path):
    # docx_grouped_images: one w:drawing holding two a:blip refs is TWO images.
    doc = docx.Document()
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(io.BytesIO(_png()))
    # Splice the second blip into the FIRST drawing to simulate a group.
    run2 = p.add_run()
    run2.add_picture(io.BytesIO(_png(6, 6)))
    drawings = p._p.findall(".//" + qn("w:drawing"))
    blip2 = drawings[1].find(".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip")
    drawings[0].append(blip2)
    run2._r.getparent().remove(run2._r)

    rec = _extract(doc, tmp_path)
    assert len(rec.figures) == 2  # one drawing, two blips → two figures
    assert rec.vector_shape_count == 0


# --- heading enumeration stripping --------------------------------------------------
def test_heading_enumeration_stripped_symmetrically():
    from docx_parse_eval.adapters.docling_adapter import extract_from_dict
    from docx_parse_eval.normalize import strip_enumeration

    assert strip_enumeration("1.2 Scope") == "Scope"
    assert strip_enumeration("2.1.3. Tolerances") == "Tolerances"
    assert strip_enumeration("Scope") == "Scope"

    rec = extract_from_dict(_docling_doc(
        texts=[{"self_ref": "#/texts/0", "label": "section_header",
                "text": "1.1 Section 1.1", "level": 2}],
        body_children=["#/texts/0"],
    ))
    assert rec.headings[0].text == "Section 1.1"
    assert "1.1 Section" not in (rec.full_text_normalized or "")


# --- docling inline groups are one paragraph ----------------------------------------
def test_docling_inline_group_is_one_paragraph():
    # An inline group = one source paragraph split at formatting boundaries;
    # unit_test_formatting turned 9 blocks into 23 without this.
    from docx_parse_eval.adapters.docling_adapter import extract_from_dict

    rec = extract_from_dict(_docling_doc(
        texts=[
            {"self_ref": "#/texts/0", "label": "text", "text": "plain"},
            {"self_ref": "#/texts/1", "label": "text", "text": "bold"},
            {"self_ref": "#/texts/2", "label": "text", "text": "italic"},
        ],
        groups=[{"self_ref": "#/groups/0", "label": "inline", "name": "group",
                 "children": [{"$ref": f"#/texts/{i}"} for i in range(3)]}],
        body_children=["#/groups/0"],
    ))
    assert rec.element_sequence == ["paragraph"]
    assert rec.full_text_normalized == "plain bold italic"


# --- TEDS quality tier ---------------------------------------------------------------
def test_teds_identical_tables_score_one():
    from docx_parse_eval.schema import TableCellRecord
    from docx_parse_eval.teds import teds

    cells = [TableCellRecord(row=r, col=c, text=f"r{r}c{c}") for r in range(3) for c in range(3)]
    assert teds(cells, cells) == 1.0
    assert teds(cells, cells, structure_only=True) == 1.0


def test_teds_separates_text_damage_from_structure_damage():
    from docx_parse_eval.schema import TableCellRecord
    from docx_parse_eval.teds import teds

    gold = [TableCellRecord(row=r, col=c, text=f"value {r}{c}") for r in range(3) for c in range(3)]
    # Text damage: same grid, one garbled cell → full TEDS drops, Struct stays 1.
    garbled = [c.model_copy() for c in gold]
    garbled[4].text = "zzzzzzzz"
    assert teds(gold, garbled) < 1.0
    assert teds(gold, garbled, structure_only=True) == 1.0
    # Structure damage: merge two cells → BOTH drop.
    merged = [c.model_copy() for c in gold[:-1]]
    merged[-1].col_span = 2
    assert teds(gold, merged) < 1.0
    assert teds(gold, merged, structure_only=True) < 1.0


def test_teds_small_typo_stays_above_threshold_garble_falls_below():
    # Boundary pin for TEDS_THRESHOLD (calibration record): a one-char typo in
    # one cell of a 3×3 table is formatting-scale noise; a garbled cell is not.
    from docx_parse_eval.schema import TableCellRecord
    from docx_parse_eval.teds import teds

    gold = [TableCellRecord(row=r, col=c, text=f"value {r}{c}") for r in range(3) for c in range(3)]
    typo = [c.model_copy() for c in gold]
    typo[4].text = typo[4].text.replace("value", "velue")
    assert teds(gold, typo) >= C.TEDS_THRESHOLD
    garbled = [c.model_copy() for c in gold]
    garbled[4].text = "garbled-beyond-recognition-xxxxx"
    assert teds(gold, garbled) < C.TEDS_THRESHOLD


def test_teds_emitted_by_comparator_on_adapter_records(tmp_path):
    # End-to-end: OOXML extraction of the synthetic docx vs its gold must score
    # TEDS == 1.0 on both tables (incl. the merged-cell table B).
    gold = read_record(FIXTURES / "synthetic.gold.json")
    ox = ooxml_reference.extract(FIXTURES / "synthetic.docx")
    results = {r.metric: r for r in compare(gold, ox)}
    assert results["table_teds"].ratio_or_score == 1.0
    assert results["table_teds_struct"].ratio_or_score == 1.0
    assert not results["table_teds"].flag


def test_identifier_overlap_is_multiset():
    gold = read_record(FIXTURES / "synthetic.gold.json")
    pred = gold.model_copy(deep=True)
    # Duplicate an existing token 5× on the gold side, keep prediction as-is:
    # a set comparison sees identical sets; the multiset must flag the loss.
    gold.identifier_tokens = sorted(gold.identifier_tokens + [gold.identifier_tokens[0]] * 5)
    results = compare(gold, pred)
    assert "identifier_token_overlap" in fired_flags(results)
