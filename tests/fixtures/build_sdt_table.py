"""Minimal synthetic fixture distilling the 'table nested in a content control'
gap found on the real public corpus (docx_rich_tables_01): a `w:tbl` wrapped in
a `w:sdt`/`w:sdtContent` block is invisible to a top-level body walk.

Authors `sdt_table.docx`: one intro paragraph + a 2×2 table that is moved inside
a content control. A correct extractor must descend into `w:sdtContent` and still
find the table (and its cell text). Built per R8/§6.3 — fix Script 1 against this,
not against the 100-page real docs.

Run: guix shell -m evaluation/manifest.scm -- python3 evaluation/tests/fixtures/build_sdt_table.py
"""

from __future__ import annotations

import io
import sys
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
DOCX_PATH = HERE / "sdt_table.docx"


def build() -> None:
    doc = docx.Document()
    doc.add_paragraph("Intro paragraph before the wrapped table.")

    t = doc.add_table(rows=2, cols=2)
    cells = [["Name", "Value"], ["Width", "120 mm"]]
    for r in range(2):
        for c in range(2):
            t.cell(r, c).text = cells[r][c]

    # Wrap the table's <w:tbl> inside <w:sdt><w:sdtContent>…</w:sdtContent></w:sdt>.
    tbl = t._tbl
    body = tbl.getparent()
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    body.insert(list(body).index(tbl), sdt)
    sdt.append(OxmlElement("w:sdtPr"))
    sdt.append(sdt_content)
    body.remove(tbl)
    sdt_content.append(tbl)

    buf = io.BytesIO()
    doc.save(buf)
    # deterministic zip (zeroed timestamps)
    src = zipfile.ZipFile(io.BytesIO(buf.getvalue()))
    with zipfile.ZipFile(DOCX_PATH, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in sorted(src.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            dst.writestr(info, src.read(name))
    print(f"wrote {DOCX_PATH.name} (table wrapped in w:sdt)")


if __name__ == "__main__":
    build()
