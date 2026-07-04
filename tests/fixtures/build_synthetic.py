"""Phase 1.b/1.c — author the synthetic `.docx` AND emit its gold record.

The `.docx` and its gold `EvaluationRecord` are two outputs of ONE construction
pass: the script *knows* what it built (counts, levels, cell texts), so the gold
is the known answer **by construction** — no human blessing, no circularity
(R9, spec §6.2). Text fields are run through the shared `normalize.py` so the
gold matches exactly what adapters will later produce.

Run:  guix shell -m evaluation/manifest.scm -- python3 evaluation/tests/fixtures/build_synthetic.py

Emits, next to this file:
  synthetic.docx        — byte-stable (zip timestamps zeroed) so git doesn't churn
  synthetic.gold.json   — gold EvaluationRecord, producer="fixture-construction"
"""

from __future__ import annotations

import io
import struct
import sys
import zlib
from pathlib import Path
import binascii
import zipfile

# Make the package importable when run as a plain script.
sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

import docx  # noqa: E402  (python-docx)
from docx.oxml import OxmlElement  # noqa: E402

from docx_parse_eval.io import sha256_file, write_record  # noqa: E402
from docx_parse_eval.normalize import (  # noqa: E402
    char_count_normalized,
    extract_identifier_tokens,
    normalize_text,
    word_count,
)
from docx_parse_eval.schema import (  # noqa: E402
    EvaluationRecord,
    FigureRecord,
    HeadingRecord,
    ListRecord,
    TableCellRecord,
    TableRecord,
)

HERE = Path(__file__).resolve().parent
DOCX_PATH = HERE / "synthetic.docx"
GOLD_PATH = HERE / "synthetic.gold.json"

PRODUCER = "fixture-construction"
PRODUCER_VERSION = "1"
DOC_ID = "synthetic-001"
TITLE = "Synthetic Warehouse Specification"

def _make_png(width: int = 4, height: int = 4) -> bytes:
    """A valid, deterministic grayscale PNG (no PIL dependency)."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", binascii.crc32(tag + data) & 0xFFFFFFFF)
        )

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 0, 0, 0, 0)  # 8-bit grayscale
    raw = b"".join(b"\x00" + b"\x80" * width for _ in range(height))  # filter 0 + gray rows
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw, 9))
        + chunk(b"IEND", b"")
    )


# Deterministic image payload for embedded figures.
_PNG = _make_png()

# Candidate special characters we deliberately seed into the body text.
_SPECIAL_CANDIDATES = ["°", "×", "±", "µ", "%"]

# --- Table contents (explicit → gold cell_text_length is a known answer) -----
TABLE_A_ROWS = [
    ["Item", "Qty", "Part No"],          # header row
    ["Conveyor", "2", "PN-12345"],
    ["Pallet rack", "10", "PN-67890"],
    ["Scanner", "5", "SC-001"],
]
# Table B: 3x3 with row 0 merged into a single spanning title cell.
TABLE_B_TITLE = "Equipment Tolerances"
TABLE_B_BODY = [
    ["Clearance", "±0.5°", "±1.0°"],
    ["Voltage", "24 V", "48 V"],
]


def _cell_text_length(cells: list[str]) -> int:
    return sum(char_count_normalized(c) for c in cells)


def build() -> EvaluationRecord:
    doc = docx.Document()

    element_sequence: list[str] = []
    text_blocks: list[str] = []
    headings: list[HeadingRecord] = []
    tables: list[TableRecord] = []
    figures: list[FigureRecord] = []
    lists: list[ListRecord] = []
    position = 0

    def add_heading(text: str, level: int) -> None:
        nonlocal position
        doc.add_heading(text, level=level)
        headings.append(HeadingRecord(text=normalize_text(text), level=level, position=position))
        element_sequence.append("heading")
        text_blocks.append(text)
        position += 1

    def add_para(text: str) -> None:
        nonlocal position
        doc.add_paragraph(text)
        element_sequence.append("paragraph")
        text_blocks.append(text)
        position += 1

    def add_list(items: list[str], ordered: bool) -> None:
        nonlocal position
        style = "List Number" if ordered else "List Bullet"
        for it in items:
            doc.add_paragraph(it, style=style)
            text_blocks.append(it)
        lists.append(
            ListRecord(
                list_id=f"l{len(lists)}",
                position=position,
                n_items=len(items),
                is_ordered=ordered,
            )
        )
        element_sequence.append("list")
        position += 1

    def add_table(rows: list[list[str]], has_header: bool, merge_first_row: bool) -> None:
        nonlocal position
        n_rows, n_cols = len(rows), len(rows[0])
        t = doc.add_table(rows=n_rows, cols=n_cols)
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                t.cell(r, c).text = val
        if has_header:
            # Encode the header row in OOXML (<w:trPr><w:tblHeader/>) so an
            # honest extractor can *derive* has_header rather than guess.
            tr_pr = t.rows[0]._tr.get_or_add_trPr()
            tr_pr.append(OxmlElement("w:tblHeader"))
        flat_cells = [v for row in rows for v in row]
        if merge_first_row:
            t.cell(0, 0).merge(t.cell(0, n_cols - 1))
            # distinct cells: 1 spanning + the remaining full rows
            cell_count = 1 + (n_rows - 1) * n_cols
            distinct_texts = [rows[0][0]] + [v for row in rows[1:] for v in row]
            cells = [TableCellRecord(row=0, col=0, col_span=n_cols, text=normalize_text(rows[0][0]))] + [
                TableCellRecord(row=r, col=c, text=normalize_text(val))
                for r, row in enumerate(rows[1:], start=1)
                for c, val in enumerate(row)
            ]
        else:
            cell_count = n_rows * n_cols
            distinct_texts = flat_cells
            cells = [
                TableCellRecord(row=r, col=c, text=normalize_text(val))
                for r, row in enumerate(rows)
                for c, val in enumerate(row)
            ]
        tables.append(
            TableRecord(
                table_id=f"t{len(tables)}",
                position=position,
                n_rows=n_rows,
                n_cols=n_cols,
                cell_count=cell_count,
                has_header=has_header,
                cell_text_length=_cell_text_length(distinct_texts),
                cells=cells,
            )
        )
        for v in distinct_texts:
            text_blocks.append(v)
        element_sequence.append("table")
        position += 1

    def add_figure(caption: str | None) -> None:
        nonlocal position
        doc.add_picture(io.BytesIO(_PNG))
        figures.append(
            FigureRecord(
                figure_id=f"f{len(figures)}",
                position=position,
                caption_text=normalize_text(caption) if caption else "",
                has_caption=caption is not None,
            )
        )
        element_sequence.append("figure")
        position += 1
        if caption is not None:
            doc.add_paragraph(caption, style="Caption")
            element_sequence.append("caption")
            text_blocks.append(caption)
            position += 1

    # --- compose the document (reading order) --------------------------------
    add_heading("Warehouse Technical Specification", 1)              # H1
    add_para("This document specifies the warehouse system at a high level.")
    add_heading("Scope", 2)                                          # H2
    add_para("Covers conveyors, racking and scanning equipment.")
    add_heading("Dimensions", 3)                                     # H3
    add_para("Each module measures 120 mm wide and draws 24 V; see PN-12345.")
    add_heading("Tolerances", 3)                                     # H3
    add_para("Permitted deviation is ±0.5° across 100% of units (µ-class).")
    add_list(["Receive goods", "Store goods", "Pick goods"], ordered=True)
    add_heading("Layout", 2)                                         # H2
    add_table(TABLE_A_ROWS, has_header=True, merge_first_row=False)
    add_figure("Figure 1: Conveyor layout overview.")
    add_heading("Equipment", 2)                                      # H2
    add_table([[TABLE_B_TITLE, "", ""], *TABLE_B_BODY], has_header=False, merge_first_row=True)
    add_figure("Figure 2: Rack elevation.")
    add_figure(None)  # uncaptioned figure
    add_heading("References", 1)                                     # H1
    add_para("Refer to standard EN-1234 for compliance details.")

    # --- persist the docx (deterministic bytes) ------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    _write_deterministic_zip(buf.getvalue(), DOCX_PATH)

    # --- derive text-level gold fields from the same strings -----------------
    full_text = normalize_text(" ".join(text_blocks))
    special = [c for c in _SPECIAL_CANDIDATES if c in full_text]

    return EvaluationRecord(
        doc_id=DOC_ID,
        title=TITLE,
        source_path=str(DOCX_PATH.relative_to(HERE.parents[2])),
        source_sha256=sha256_file(DOCX_PATH),
        producer=PRODUCER,
        producer_version=PRODUCER_VERSION,
        word_count=sum(word_count(b) for b in text_blocks),
        char_count_normalized=char_count_normalized(full_text),
        full_text_normalized=full_text,
        tables=tables,
        figures=figures,
        headings=headings,
        lists=lists,
        hyperlink_count=0,
        footnote_count=0,
        endnote_count=0,
        vector_shape_count=0,  # all synthetic figures are raster images
        element_sequence=element_sequence,
        identifier_tokens=extract_identifier_tokens(full_text),
        special_chars=special,
    )


def _write_deterministic_zip(raw: bytes, out: Path) -> None:
    """Rewrite a docx (zip) with zeroed timestamps so regeneration is byte-stable
    and git doesn't churn on identical content."""
    src = zipfile.ZipFile(io.BytesIO(raw))
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as dst:
        for name in sorted(src.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            dst.writestr(info, src.read(name))


def main() -> None:
    record = build()
    write_record(record, GOLD_PATH)
    print(f"wrote {DOCX_PATH.name} (sha256={record.source_sha256[:12]}…)")
    print(f"wrote {GOLD_PATH.name}: {len(record.tables)} tables, "
          f"{len(record.figures)} figures, {len(record.headings)} headings, "
          f"{len(record.lists)} lists, seq_len={len(record.element_sequence)}")


if __name__ == "__main__":
    main()
